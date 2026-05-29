# -*- coding: utf-8 -*-
import os
import re
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from playwright.sync_api import sync_playwright

def set_cell_background(cell, fill_hex):
    """Sets background color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        m = OxmlElement(margin)
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Applies neat subtle borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx_from_markdown(md_path, docx_path):
    print("Generating DOCX...")
    doc = Document()
    
    # Configure page setup (A4 margins)
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69) # A4 height

    # Base colors
    primary_color = RGBColor(0x3B, 0x45, 0xC4)  # #3b45c4
    dark_gray = RGBColor(0x2A, 0x2A, 0x2A)      # #2a2a2a
    body_color = RGBColor(0x1C, 0x1C, 0x1C)     # #1c1c1c
    code_color = RGBColor(0x2A, 0x2A, 0x8A)     # #2a2a8a

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_headers = []
    table_rows = []
    
    in_code_block = False
    code_lines = []

    def process_run_text(p, text, is_italic=False, custom_size=10.5):
        """Helper to process simple markdown formatting inside a paragraph."""
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
            
            # Bold
            if part.startswith('**') and part.endswith('**'):
                inner = part[2:-2]
                run = p.add_run(inner)
                run.bold = True
                run.italic = is_italic
                run.font.name = 'Arial'
                run.font.size = Pt(custom_size)
                run.font.color.rgb = body_color
            # Inline code
            elif part.startswith('`') and part.endswith('`'):
                inner = part[1:-1]
                run = p.add_run(inner)
                run.italic = is_italic
                run.font.name = 'Consolas'
                run.font.size = Pt(custom_size - 1.0)
                run.font.color.rgb = code_color
            else:
                run = p.add_run(part)
                run.bold = False
                run.italic = is_italic
                run.font.name = 'Arial'
                run.font.size = Pt(custom_size)
                run.font.color.rgb = body_color

    def add_styled_paragraph(text, space_before=0, space_after=6, bullet=False, quote=False):
        if bullet:
            p = doc.add_paragraph(style='List Bullet')
        else:
            p = doc.add_paragraph()
            
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if quote:
            p.paragraph_format.left_indent = Inches(0.4)
            process_run_text(p, text, is_italic=True)
        else:
            process_run_text(p, text)
        return p

    def add_styled_heading(text, level, space_before=16, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        
        if level == 1:
            run.font.size = Pt(17)
            run.font.color.rgb = primary_color
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = primary_color
        else:
            run.font.size = Pt(11.5)
            run.font.color.rgb = dark_gray
        return p

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if not table_headers and not table_rows:
            return
            
        # Create table
        cols_count = len(table_headers) if table_headers else (len(table_rows[0]) if table_rows else 1)
        table = doc.add_table(rows=0, cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table borders
        set_table_borders(table)

        # Header Row
        if table_headers:
            row = table.add_row()
            # Set repeat header and keep with next
            trPr = row._tr.get_or_add_trPr()
            trPr.append(OxmlElement('w:tblHeader'))
            trPr.append(OxmlElement('w:cantSplit'))
            
            for i, header_text in enumerate(table_headers):
                cell = row.cells[i]
                set_cell_background(cell, "5C66DE") # Header background #5c66de
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                
                # Write header text
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                run = p.add_run(header_text.strip())
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text

        # Data Rows
        for row_idx, row_data in enumerate(table_rows):
            row = table.add_row()
            # Keep row together
            trPr = row._tr.get_or_add_trPr()
            trPr.append(OxmlElement('w:cantSplit'))
            
            is_even = (row_idx % 2 == 1)
            bg_color = "F3F4FB" if is_even else "FFFFFF"
            
            for i in range(min(cols_count, len(row_data))):
                cell = row.cells[i]
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=90, bottom=90, left=150, right=150)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                
                process_run_text(p, row_data[i].strip(), custom_size=9.0)
                
        # Empty space after table
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        
        # Reset table data
        table_headers = []
        table_rows = []
        in_table = False

    def flush_code_block():
        nonlocal in_code_block, code_lines
        if not code_lines:
            return
        
        # Create a table with 1 column/1 row for shaded code background
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F3F4FB")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        # Set border
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D6EA"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D6EA"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D3D6EA"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D3D6EA"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        first = True
        for line in code_lines:
            if not first:
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)
            run.font.color.rgb = body_color
            first = False
            
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        
        code_lines = []
        in_code_block = False

    for line_raw in lines:
        line = line_raw.rstrip('\n')
        line_stripped = line.strip()

        # Handle fenced code block
        if line_stripped.startswith('```'):
            if in_code_block:
                flush_code_block()
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue

        # Handle Markdown Tables
        if line_stripped.startswith('|'):
            if not in_table:
                if in_code_block:
                    flush_code_block()
                in_table = True
                # Parse header
                cols = [c.strip() for c in line_stripped.split('|')[1:-1]]
                table_headers = cols
            else:
                # Check for divider line like |---|---|
                if re.match(r'^\|[\s:-|-]*\|$', line_stripped) or '---' in line_stripped:
                    continue
                cols = [c.strip() for c in line_stripped.split('|')[1:-1]]
                table_rows.append(cols)
            continue
        elif in_table:
            flush_table()

        # Handle Empty Lines
        if not line_stripped:
            continue

        # Horizontal rule
        if line_stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="DCDCE4"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            continue

        # Headings
        match = re.match(r'^(#{1,6})\s+(.*)$', line_stripped)
        if match:
            level = len(match.group(1))
            heading_text = match.group(2)
            add_styled_heading(heading_text, level)
            continue

        # Blockquotes
        if line_stripped.startswith('>'):
            quote_text = line_stripped[1:].strip()
            add_styled_paragraph(quote_text, quote=True)
            continue

        # Bullet list items
        if line_stripped.startswith(('-', '*', '+')) and not line_stripped.startswith('---'):
            list_text = re.sub(r'^[\-\*\+]\s+', '', line_stripped)
            add_styled_paragraph(list_text, bullet=True)
            continue

        # Numbered list items
        num_match = re.match(r'^(\d+)\.\s+(.*)$', line_stripped)
        if num_match:
            # We will render simple paragraph with bold numbers for better reliability in styles
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            run_num = p.add_run(f"{num_match.group(1)}. ")
            run_num.bold = True
            run_num.font.name = 'Arial'
            run_num.font.size = Pt(10.5)
            run_num.font.color.rgb = body_color
            
            process_run_text(p, num_match.group(2))
            continue

        # Regular Paragraph
        add_styled_paragraph(line_stripped)

    # Save documents
    if in_table:
        flush_table()
    if in_code_block:
        flush_code_block()

    doc.save(docx_path)
    print(f"DOCX created successfully at {docx_path}")

def build_html_from_markdown(md_path, html_path):
    print("Generating HTML...")
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Convert markdown to HTML
    html_body = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])

    # Post-process HTML to wrap tables in `<div class="tw">` and add `border="1"`
    def wrap_table(match):
        table_content = match.group(0)
        if 'border=' not in table_content:
            table_content = table_content.replace('<table', '<table border="1"')
        return f'<div class="tw">{table_content}</div>'

    html_body = re.sub(r'<table.*?>.*?</table>', wrap_table, html_body, flags=re.DOTALL)

    style_header = """<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>Режимы Мыйзамчы</title><style>
body{font-family:'Segoe UI',Arial,sans-serif;font-size:15px;line-height:1.55;color:#1c1c1c;max-width:840px;margin:0 auto;padding:18px;}
h1{font-size:23px;color:#3b45c4;border-bottom:2px solid #5C66DE;padding-bottom:6px;margin-top:8px;}
h2{font-size:19px;color:#3b45c4;margin-top:26px;border-bottom:1px solid #dfe1f0;padding-bottom:3px;}
h3{font-size:16px;color:#2a2a2a;margin-top:18px;}
p{margin:8px 0;}
.tw{overflow-x:auto;margin:12px 0;}
table{border-collapse:collapse;width:100%;font-size:12.5px;}
th,td{border:1px solid #c4c4d4;padding:5px 7px;text-align:left;vertical-align:top;}
th{background:#5C66DE;color:#ffffff;font-weight:600;}
tr:nth-child(even) td{background:#f3f4fb;}
code{background:#eceef7;padding:1px 4px;border-radius:3px;font-family:Consolas,'Courier New',monospace;font-size:90%;color:#2a2a8a;}
pre{background:#f3f4fb;border:1px solid #d3d6ea;padding:11px 13px;border-radius:6px;overflow-x:auto;font-size:12.5px;line-height:1.4;}
pre code{background:none;color:#1c1c1c;padding:0;}
blockquote{border-left:3px solid #5C66DE;margin:10px 0;padding:6px 12px;background:#f3f4fb;color:#444;}
hr{border:none;border-top:1px solid #dcdce4;margin:22px 0;}
</style></head><body>"""

    full_html = style_header + html_body + "</body></html>"

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    print(f"HTML created successfully at {html_path}")

def build_pdf_from_html(html_path, pdf_path):
    print("Generating PDF via Playwright...")
    abs_html_path = os.path.abspath(html_path)
    file_url = f"file:///{abs_html_path.replace(os.sep, '/')}"
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(file_url)
        page.wait_for_load_state("networkidle")
        page.pdf(
            path=pdf_path,
            format="A4",
            margin={"top": "20mm", "bottom": "20mm", "left": "20mm", "right": "20mm"},
            print_background=True
        )
        browser.close()
    print(f"PDF created successfully at {pdf_path}")

if __name__ == '__main__':
    md_path = 'РЕЖИМЫ.md'
    docx_path = 'РЕЖИМЫ.docx'
    html_path = 'РЕЖИМЫ.html'
    pdf_path = 'РЕЖИМЫ.pdf'
    
    build_html_from_markdown(md_path, html_path)
    build_docx_from_markdown(md_path, docx_path)
    try:
        build_pdf_from_html(html_path, pdf_path)
    except Exception as e:
        print(f"PDF generation warning (skipped): {e}")
    print("All documentation files compiled successfully!")
